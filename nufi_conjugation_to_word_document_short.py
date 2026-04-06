import random
import docx
from docx.enum import text as _text, style as _style, table as _table
import json
import requests
import re
import time as _time

# Path and file settings
file_name = 'conjugaison_nufi__short2.docx'
path_verbs = r"G:\My Drive\Data_Science\DataSets\data_verb.json"
nb_table_per_page = 4


# Requirements for verb conjugation
requirements = {
    'simple': {
        1: {
            1: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            2: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            3: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,}
        },
        2: {
            1: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            2: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            3: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,}
        },
    },
    'compose': {
        1: {
            1: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            2: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            3: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,}
        },
        2: {
            1: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            2: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,},
            3: {'mb': 5, 'nd': 5, 'nj': 5, 'ng': 5, 'ngw': 5,}
        },
    },
}


# We set the requirements
requirements_short = {
    # category
    'simple': {
        # group
        1: {
            # x_syllabes: number of syllabes
            1: {
                # constraints: need
                 'nd': 3,
            },

            2: {'mb': 2, 'nd': 2, 'nj': 2, 'ng': 2, 'ngw': 2,},
           
        },
        
    },
    
}
# Time groups mapping to conjugation tenses
time_groups = {
    'Présent': {
        0: 'Présent',
        10: 'Impératif'
    },
    'Passé Récent': {
        1: 'Passé Composé',
        2: 'Passé Récent',
    },
    'Passé': {
        3: 'Passé d\'hier',
        4: 'Passé (>= 2 jours)',
        6: 'Passé Lointain',
    },
    'Futur Proche': {
        7: 'Futur Proche',
    },
    'Futur Lointain': {
        8: 'Futur Lointain',
    },
    'Temps Habituel': {
        9: 'Présent Habituel',
        5: "Passé Habituel"
    },
}

# Map each tense group to a specific table style
tense_style_mapping = {
    'Présent': 'Light List Accent 1',
    'Passé Récent': 'Light List Accent 2',
    'Passé': 'Light List Accent 3',
    'Futur Proche': 'Light Grid',
    'Futur Lointain': 'Light List Accent 4',
    'Temps Habituel': 'Light Grid Accent 5'
}

def insert_table(document: docx.document.Document, headers: list, data: list, title: str, table_style: str) -> None:
    heading = document.add_heading(title, level=1)
    heading.alignment = _text.WD_ALIGN_PARAGRAPH.CENTER
    heading.bold = False

    rows = len(data)
    columns = len(headers)

    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    table = document.add_table(rows+1, columns)
    table.alignment = _table.WD_TABLE_ALIGNMENT.CENTER

    # Adding style to a table
    table.style = table_style
    table.allow_autofit = True
    table.autofit = True

    # add the header rows
    for j in range(columns):
        cell = table.cell(0, j)
        pararagraph = cell.paragraphs[0]
        pararagraph.alignment = _text.WD_ALIGN_PARAGRAPH.CENTER
        pararagraph.add_run(headers[j])
        pararagraph.style='Nufi'
        pararagraph.bold = False

    # add the rest of the data
    for i in range(rows):
        for j in range(columns):
            cell = table.cell(i+1, j)
            pararagraph = cell.paragraphs[0]
            pararagraph.style = 'Nufi'
            pararagraph.alignment = _text.WD_ALIGN_PARAGRAPH.CENTER
            pararagraph.add_run(data[i][j]).bold = False

def get_available_conjugation_times() -> list:
    data = requests.get(f"https://resulam.com/nufi-conjugator/?api=1").json()
    return data.get('available_times', [])

from bs4 import BeautifulSoup

# response = requests.get(f"http://127.0.0.1/nufi-conjugator/?verb={verb}&time={conj_time}")
# print(response.text)  # This will show the raw response

# with open("response_conjugaison.txt","w") as f:
#     f.write(response.text)

def get_conjugaison(verb: str) -> list:

    # verb = "nduu2'"
    # conj_time = "conj_mode_imperatif"
    # conj_time = "conj_present1"
    conj = {}
    for conj_time in available_conjugation_times:
        data = requests.get(f"https://resulam.com/nufi-conjugator/?api=1&verb={verb}&time={conj_time}").json()
        
        if not "verb" in data:
            return []

        conj['translated_verb'] = data['translated_verb']
        conj['verb'] = data['verb']
        conj[conj_time] = data['data'][0]['data']
    return conj

# def get_conjugaison(verb):

def get_conjugaison_localhost(verb):

    # data = requests.get(f"http://127.0.0.1/nufi-conjugator/?verb={verb}&time={conj_time}").json()
    # verb = "ndhíshʉ̄ɑ̄'"
    response = requests.get(f"http://127.0.0.1/nufi-conjugator/?verb={verb}")
    # Parse the HTML
    soup = BeautifulSoup(response.text, 'html.parser')

    # Find all div elements with the class 'card-header'
    card_headers = soup.find_all('div', class_='card-header')
    card_headers = card_headers[1:]  # If you need to skip the first element

    # Extract text directly from each card header without re-parsing
    conj_tenses = [div.get_text() for div in card_headers]
    conj_tenses_cleaned = [text.replace('\u200b', '') for text in conj_tenses]

    # target_tense = conj_tenses[3]

    conj_tenses_dict = {"Ntìé'è lè (présent progressif, maintenant)": 'conj_present1',
    "Tōh mēndɑ̀' (passé récent ou passé composé)": 'conj_passee1',
    "Ntìè' ē tōh ngwǎ’ nàm ntōhō (passé récent)": 'conj_passee2',
    "Ntìè' ē tōh wāhà\u200b (le passé d'hier)": 'conj_passee3',
    "Ntìè' ē tōh līē' pʉ́ɑ\u200b (passé supérieur ou égal  à deux jours)": 'conj_passee4',
    "Ntìè' ē tōh ndìàndìà (passé habituel, imparfait de l'indicatif)": 'conj_passee5',
    "Ntìè' ē tōh tɑ̀ kwàŋ\u200b\u200b (le passé lointain)": 'conj_passee6',
    "Ntìē'sɑ̀' mêndɑ̀'\u200b (le futur proche)": 'conj_futur1',
    "Ntìē'sɑ̀' sʉ̀sʉ̀ɑ̀\u200b (le futur lointain)": 'conj_futur2',
    "Ntìè' ē ghʉ̀ ndìàndìà\u200b (le présent d'habitude)": 'conj_present2',
    "Ghə̀ə̄ngʉ̀'\u200b (mode imperatif)": 'conj_mode_imperatif'}

    # Ntìè' ē tōh wāhà​ (le passé d'hier)

    # Parse the HTML content
    
    # Find all div elements with the class 'card-header'
    card_headers2 = soup.find_all('div', class_='card-header')
    conjugations_dict = {}
    conjugations_dict["verb"] = verb
    for target_tense in conj_tenses:
            
        for header in card_headers2:
            # Check if the header text matches the target header
            if header.get_text().strip() == target_tense:
                # Find the next sibling of the header which is likely the conjugation block
                conjugation_block = header.find_next_sibling('div', class_='card-block')
                # Extract conjugations from the table within this block
                if conjugation_block:
                    conjugations = []
                    # Find all rows in the table and extract the text
                    rows = conjugation_block.find_all('tr')
                    conjugations = [[td.get_text(strip=True) for td in row.find_all('td')] for row in rows]
                    
                    conj_tense_nick_name = conj_tenses_dict[target_tense]
                    conjugations_dict[conj_tense_nick_name] = conjugations
                
    return conjugations_dict
    
# res = get_conjugaison(verb)
# len(res.keys())
# res2 = get_conjugaison_localhost(verb)
# len(res2.keys())

print('Loading data...')
with open(path_verbs, 'r', encoding='utf-8') as f:
    data_verb = json.load(f)   
    
    data_verb_dict = {i["verb"]:i["translated_verb"] for i in data_verb}
    data_verb = [i["verb"] for i in data_verb]

print('Downloading available conjugation times...')
available_conjugation_times = get_available_conjugation_times()

if not available_conjugation_times:
    raise Exception("No conjugation time availables on the server.")

print('Configuration document...')
document = docx.Document()
style = document.styles.add_style('Nufi', _style.WD_STYLE_TYPE.PARAGRAPH)
style.font.size = docx.shared.Pt(12)
style.font.name = 'Charis-SIL'
style.font.bold = False

print('Building documents...')
for category, groups in requirements_short.items():
    document.add_page_break()
    document.add_heading(category, level=0)

    for group, x_syllabes in groups.items():
        document.add_page_break()
        document.add_heading(f'Group {group}', level=0)

        for nb_syllabe, needs in x_syllabes.items():
            document.add_page_break()
            document.add_heading(f'{nb_syllabe} syllabe(s)', level=0)

            for constraint in needs:
                document.add_page_break()
                document.add_heading(f'Verb start with {constraint}', level=0)

                print(f'\tcategory: {category}')
                print(f'\tgroup: {group}')
                print(f'\tnb_syllabe: {nb_syllabe}')
                print(f'\tconstraint: {constraint}')

                for verb in data_verb:
                    if needs[constraint] <= 0:
                        break

                    nb_table = 0

                    if not verb.startswith(constraint):
                        continue

                    data_verb.remove(verb)

                    print(f'\t\tGetting conjugaison of {verb}...')

                    delay = 1
                    while True:
                        try:
                            _time.sleep(delay)
                            conj = get_conjugaison(verb)
                            break
                        except Exception as e:
                            print("Error: ", e)
                            delay += 1
                            print(f'\t\tRetry with a delay of {delay}s...')
                            _time.sleep(delay)

                    if not conj:
                        continue

                    needs[constraint] -= 1

                    print(f'\t\tBuilding conjugaison of {verb}...')

                    try:
                            
                        for title, time_group in time_groups.items():
                            headers = []
                            data = []

                            for time, subtitle in time_group.items():
                                time = list(available_conjugation_times.keys())[time]
                                if len(conj[time][0]) > 1:
                                    headers += [f'{subtitle} {i+1}' for i in range(len(conj[time][0]))]
                                else:
                                    headers.append(subtitle)

                                if data == []:
                                    data = conj[time]
                                else:
                                    for i in range(len(conj[time])):
                                        data[i] += conj[time][i]

                            if nb_table % nb_table_per_page == 0:
                                document.add_page_break()

                            # Get the style for the current tense group
                            table_style = tense_style_mapping.get(title, 'Light Grid Accent 5')
                            insert_table(document, headers, data, f"{title} ({conj['verb']}: {data_verb_dict[verb]})", table_style)
                            nb_table += 1
                    except:
                        pass 
            print('\tSaving...')
            document.save(file_name)

print("Finish")
